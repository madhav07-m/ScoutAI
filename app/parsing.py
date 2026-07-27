"""
Phase 1 — Document ingestion & parsing.

Extracts raw text from PDF / DOCX resumes and JDs, then cleans it up.

Known gotcha: multi-column resumes extract out of reading order with
PyMuPDF's default text extraction. We mitigate this by extracting text
in "blocks" mode and sorting blocks top-to-bottom, left-to-right, which
handles most single- and two-column templates reasonably well. For
resumes that STILL extract poorly (heavily graphic/templated layouts,
tables inside tables, text boxes, or scanned/image-based PDFs), this
module now has a staged fallback instead of silently returning near-
empty text:

  1. PyMuPDF blocks-mode extraction (fast, handles most layouts).
  2. If that looks too thin (see is_low_extraction): retry with
     pdfplumber, which has a different (often better) table/layout
     model and occasionally recovers text PyMuPDF's block-sort misses.
  3. If STILL too thin: assume this is an image-based / scanned PDF
     (icon-heavy graphic resumes often rasterize their whole content)
     and fall back to OCR via pytesseract, rendering each page with
     PyMuPDF (no separate poppler/pdf2image dependency needed) and
     running Tesseract locally -- free, no API call.
  4. If OCR still comes back thin, we give up gracefully: return
     whatever text was extracted (however little) and let
     is_low_extraction() flag it downstream, and chunk_by_section()
     will bucket it into a single "General" section rather than
     erroring out. Silent misparsing becomes a visible, honestly
     labeled low-confidence result instead.

This means parse_document() may be slower for resumes that need OCR
(each fallback stage is progressively more expensive), but only pays
that cost when the cheaper stages actually failed.
"""

import re
import io
import fitz  # PyMuPDF
import docx


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF, attempting to preserve reading order."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    all_text = []

    for page in doc:
        # "blocks" gives (x0, y0, x1, y1, text, block_no, block_type)
        blocks = page.get_text("blocks")
        # Sort top-to-bottom, then left-to-right — a cheap but effective
        # fix for simple multi-column layouts.
        blocks.sort(key=lambda b: (round(b[1], 1), round(b[0], 1)))
        for b in blocks:
            text = b[4].strip()
            if text:
                all_text.append(text)

    doc.close()
    return "\n".join(all_text)


def extract_text_with_pdfplumber(file_bytes: bytes) -> str:
    """Fallback PDF extractor using pdfplumber, which uses a different
    layout/table model than PyMuPDF and sometimes recovers text from
    resumes that defeat the blocks-mode extractor above (e.g. content
    laid out in actual PDF tables, or certain text-box templates).
    Optional dependency -- returns "" if pdfplumber isn't installed
    rather than crashing the whole pipeline.
    """
    try:
        import pdfplumber
    except ImportError:
        return ""

    all_text = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                all_text.append(text.strip())
            # Also pull table cells explicitly -- pdfplumber's table
            # detection can surface content that extract_text() alone
            # misses in strongly tabular resume templates.
            for table in page.extract_tables():
                for row in table:
                    for cell in row:
                        if cell and cell.strip():
                            all_text.append(cell.strip())
    return "\n".join(all_text)


def extract_text_with_ocr(file_bytes: bytes, dpi: int = 200) -> str:
    """Last-resort OCR fallback for image-based/scanned PDFs (common
    in icon-heavy graphic resume templates that rasterize content
    instead of embedding selectable text). Renders each page to an
    image using PyMuPDF (already a dependency -- no separate poppler/
    pdf2image install needed) and runs Tesseract locally via
    pytesseract. Free, runs on-device, no API call.

    Optional dependency -- returns "" if pytesseract or the Tesseract
    binary isn't installed, rather than crashing the pipeline. On
    Windows, pytesseract additionally requires the Tesseract binary
    itself (a separate installer, not just `pip install pytesseract`)
    on PATH or pointed to via pytesseract.pytesseract.tesseract_cmd.
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""

    all_text = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        zoom = dpi / 72
        matrix = fitz.Matrix(zoom, zoom)
        for page in doc:
            pix = page.get_pixmap(matrix=matrix)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            text = pytesseract.image_to_string(img)
            if text.strip():
                all_text.append(text.strip())
        doc.close()
    except Exception:
        # Tesseract binary missing/misconfigured, corrupt PDF, etc. --
        # degrade to "no OCR text" rather than raising, so ranking can
        # still proceed with whatever the earlier stages extracted.
        return ""

    return "\n".join(all_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file, including table cells."""
    document = docx.Document(io.BytesIO(file_bytes))
    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    return "\n".join(parts)


def clean_text(raw_text: str) -> str:
    """Normalize messy extracted text.

    Handles: broken line breaks (words split mid-sentence), repeated
    whitespace, common header/footer noise (page numbers), and bullet
    character normalization.
    """
    text = raw_text

    # Collapse hyphenated line breaks: "soft-\nware" -> "software"
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)

    # Remove standalone page-number-only lines
    text = re.sub(r"\n\s*\d{1,3}\s*\n", "\n", text)

    # Normalize bullet characters to a plain dash
    text = re.sub(r"[•●▪‣◦]", "-", text)

    # Collapse 3+ newlines into 2 (paragraph-ish spacing)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Collapse repeated spaces/tabs
    text = re.sub(r"[ \t]{2,}", " ", text)

    return text.strip()


def parse_document(filename: str, file_bytes: bytes) -> str:
    """Route a file to the right extractor and clean the result.

    For PDFs, tries progressively more expensive fallbacks if the
    cheaper stage's output looks too thin (see is_low_extraction and
    the module docstring). Use parse_document_with_meta() instead if
    you need to know which stage actually produced the text (e.g. to
    surface "this was OCR'd, double-check it" in the UI).
    """
    return parse_document_with_meta(filename, file_bytes)["text"]


def parse_document_with_meta(filename: str, file_bytes: bytes) -> dict:
    """Same as parse_document, but returns {"text": ..., "method": ...}
    so callers can tell whether the result came from the fast path or
    a fallback (pdfplumber / OCR), which matters for deciding how much
    to trust a low Fit Score downstream.
    """
    lower = filename.lower()

    if lower.endswith(".pdf"):
        raw = extract_text_from_pdf(file_bytes)
        method = "pymupdf"

        if is_low_extraction(clean_text(raw)):
            plumber_raw = extract_text_with_pdfplumber(file_bytes)
            if plumber_raw and not is_low_extraction(clean_text(plumber_raw)):
                raw = plumber_raw
                method = "pdfplumber"
            else:
                # Both text-layer extractors came back thin -- likely
                # an image-based/scanned/heavily-graphic resume. Try
                # OCR as a last resort before giving up.
                ocr_raw = extract_text_with_ocr(file_bytes)
                if ocr_raw and len(ocr_raw.strip()) > len(raw.strip()):
                    raw = ocr_raw
                    method = "ocr"
                # else: keep whatever the first pass got, however
                # thin -- is_low_extraction() downstream will flag it
                # rather than this function silently pretending it's fine.

    elif lower.endswith(".docx"):
        raw = extract_text_from_docx(file_bytes)
        method = "docx"
    elif lower.endswith(".txt"):
        raw = file_bytes.decode("utf-8", errors="ignore")
        method = "txt"
    else:
        raise ValueError(f"Unsupported file type: {filename}")

    return {"text": clean_text(raw), "method": method}


def is_low_extraction(text: str, min_chars: int = 300) -> bool:
    """Heuristic flag for the known multi-column/graphic-resume gotcha:
    if extraction yielded suspiciously little text, the layout likely
    defeated the block-sorting extractor (e.g. text boxes, image-based
    sections), and a resulting low Fit Score should be read as
    "possible parsing issue" rather than "weak candidate."
    """
    return len(text.strip()) < min_chars
